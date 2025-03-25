import os
import sys
import tkinter as tk
from tkinter import filedialog
from tkinter import ttk
import pytesseract
from PIL import Image
from docx import Document
import speech_recognition as sr

# Устанавливаем кодировку UTF-8
sys.stdout.reconfigure(encoding="utf-8")

# Указываем путь к Tesseract (если требуется)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Определяем путь к рабочему столу и Word-файлу
desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
word_filename = os.path.join(desktop_path, "recognized_text.docx")

# Создаём или загружаем Word-документ
if os.path.exists(word_filename):
    doc = Document(word_filename)  # Открываем существующий файл
else:
    doc = Document()  # Создаем новый документ

# Функция для обработки изображения
def process_image():
    file_path = filedialog.askopenfilename(
        title="Выберите изображение",
        filetypes=[("Изображения", "*.png;*.jpg;*.jpeg;*.bmp;*.tiff")],
    )
    if not file_path:
        print("Файл не выбран.")
        return

    try:
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        print("Распознанный текст из изображения:\n", text)

        doc.add_heading("Распознанный текст из изображения:", level=1)
        doc.add_paragraph(text)
        doc.add_paragraph("\n" + "-" * 50 + "\n")
        doc.save(word_filename)

        print(f"Файл '{word_filename}' успешно обновлён.")

    except Exception as e:
        print("Ошибка при обработке изображения:", e)

# Функция для обработки аудио
def process_audio():
    file_path = filedialog.askopenfilename(
        title="Выберите аудиофайл", filetypes=[("Аудиофайлы", "*.wav;*.mp3;*.flac")]
    )
    if not file_path:
        print("Файл не выбран.")
        return

    recognizer = sr.Recognizer()
    try:
        with sr.AudioFile(file_path) as source:
            audio = recognizer.record(source)
            text = recognizer.recognize_google(audio, language="ru-RU")
            print("Распознанный текст из аудио:\n", text)

            doc.add_heading("Распознанный текст из аудио:", level=1)
            doc.add_paragraph(text)
            doc.add_paragraph("\n" + "-" * 50 + "\n")
            doc.save(word_filename)

            print(f"Файл '{word_filename}' успешно обновлён.")

    except Exception as e:
        print("Ошибка при обработке аудио:", e)

# Функция для создания кастомного диалогового окна
def show_custom_dialog():
    dialog = tk.Toplevel(root)
    dialog.title("Выбор")
    dialog.geometry("400x200")
    dialog.resizable(False, False)

    # Стилизация
    style = ttk.Style()
    style.configure("TButton", font=("Arial", 12), padding=10)

    # Текстовое сообщение
    label = tk.Label(dialog, text="Что хотите обработать?", font=("Arial", 12))
    label.pack(pady=15)

    # Кнопки
    btn_image = ttk.Button(dialog, text="📷 Изображение", command=lambda: [dialog.destroy(), process_image()])
    btn_image.pack(pady=5, ipadx=10)

    btn_audio = ttk.Button(dialog, text="🎙️ Аудио", command=lambda: [dialog.destroy(), process_audio()])
    btn_audio.pack(pady=5, ipadx=10)

    btn_cancel = ttk.Button(dialog, text="❌ Отмена", command=lambda: [dialog.destroy(), root.quit()])
    btn_cancel.pack(pady=10, ipadx=10)

    dialog.transient(root)  # Окно поверх основного
    dialog.grab_set()  # Захват фокуса

# Основное окно (скрываем рамку, чтобы оно не мешало)
root = tk.Tk()
root.overrideredirect(True)  # Скрываем рамку
root.geometry("1x1+0+0")  # Делаем окно минимальным
root.after(100, show_custom_dialog)  # Вызываем окно выбора

root.mainloop()  # Главный цикл Tkinter

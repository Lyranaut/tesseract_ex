import os
import tkinter as tk
from tkinter import filedialog
import speech_recognition as sr
from docx import Document

# Открываем диалоговое окно для выбора аудиофайла
root = tk.Tk()
root.withdraw()
file_path = filedialog.askopenfilename(title="Выберите аудиофайл", filetypes=[("Аудиофайлы", "*.wav;*.mp3;*.flac")])

if not file_path:
    print("Файл не выбран. Программа завершена.")
    exit()

# Создаем объект для распознавания
recognizer = sr.Recognizer()

# Определяем путь к файлу Word
desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
word_filename = os.path.join(desktop_path, "recognized_text.docx")

# Создаем или загружаем документ Word
if os.path.exists(word_filename):
    doc = Document(word_filename)  # Открываем существующий файл
else:
    doc = Document()  # Создаем новый файл

# Загружаем аудиофайл и распознаем текст
try:
    with sr.AudioFile(file_path) as source:
        audio = recognizer.record(source)
        text = recognizer.recognize_google(audio, language="ru-RU")
        print("Распознанный текст: \n", text)

        # Добавляем заголовок и текст в Word
        doc.add_heading("Распознанный текст из аудио:", level=1)
        doc.add_paragraph(text)
        doc.add_paragraph("\n" + "-" * 50 + "\n")  # Разделитель

        # Сохраняем документ
        doc.save(word_filename)
        print(f"Файл '{word_filename}' успешно создан/обновлен на рабочем столе.")

except Exception as e:
    print("Произошла ошибка при распознавании: ", e)
